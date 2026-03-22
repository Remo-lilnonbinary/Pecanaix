"""
Read CSV, Excel, PDF, Word, and plain-text files from a folder into structured alumni rows
and unstructured text blobs.

Requires: pip install openpyxl PyPDF2 python-docx
"""

from __future__ import annotations

import csv
import logging
import os
import re
from typing import Any

logger = logging.getLogger(__name__)


def _norm_header(h: str) -> str:
    """Normalize header for alias lookup (lowercase alphanumerics only)."""
    return "".join(c.lower() for c in h.strip() if c.isalnum())


def _map_header_to_field(norm: str) -> str | None:
    """Map normalized header to internal field id (including _first, _last, _full)."""
    mapping: dict[str, str] = {
        # First / last / full name
        "firstname": "_first",
        "first": "_first",
        "lastname": "_last",
        "last": "_last",
        "surname": "_last",
        "fullname": "_full",
        "name": "name_single",
        # Email
        "email": "email",
        "emailaddress": "email",
        # Graduation year
        "graduationyear": "graduation_year",
        "gradyear": "graduation_year",
        "year": "graduation_year",
        # Location
        "city": "location_city",
        "location": "location_city",
        "mailingcity": "location_city",
        # Industry
        "industry": "industry",
        # Job title
        "jobtitle": "job_title",
        "title": "job_title",
        "role": "job_title",
        "currentrole": "job_title",
        # Company
        "company": "company",
        "currentcompany": "company",
        "organization": "company",
        # Department
        "department": "department",
        "dept": "department",
    }
    return mapping.get(norm)


def _cell_str(val: Any) -> str:
    if val is None:
        return ""
    if isinstance(val, float) and val == int(val):
        return str(int(val))
    return str(val).strip()


def _parse_graduation_year(val: str) -> int | None:
    if not val:
        return None
    m = re.search(r"(\d{4})", val)
    if not m:
        return None
    try:
        y = int(m.group(1))
        if 1900 <= y <= 2100:
            return y
    except ValueError:
        pass
    return None


def _build_name(fields: dict[str, str]) -> str:
    full = (fields.get("_full") or "").strip()
    if full:
        return full
    first = (fields.get("_first") or "").strip()
    last = (fields.get("_last") or "").strip()
    if first or last:
        return " ".join(p for p in (first, last) if p)
    single = (fields.get("name_single") or "").strip()
    return single


def _row_dict_to_alumni_fields(raw_row: dict[str, Any]) -> dict[str, Any]:
    """Map arbitrary column keys to canonical alumni keys."""
    fields: dict[str, str] = {}
    for raw_key, raw_val in raw_row.items():
        if raw_key is None:
            continue
        key = str(raw_key).strip()
        if not key:
            continue
        norm = _norm_header(key)
        if not norm:
            continue
        field = _map_header_to_field(norm)
        if not field:
            continue
        val = _cell_str(raw_val)
        if not val:
            continue
        # Last write wins if duplicate mapped columns
        fields[field] = val

    name = _build_name(fields)

    out: dict[str, Any] = {
        "name": name,
        "email": fields.get("email") or "",
        "graduation_year": _parse_graduation_year(fields.get("graduation_year") or ""),
        "location_city": fields.get("location_city") or "",
        "industry": fields.get("industry") or "",
        "job_title": fields.get("job_title") or "",
        "company": fields.get("company") or "",
        "department": fields.get("department") or "",
    }
    return out


def ingest_csv(filepath: str) -> list[dict[str, Any]]:
    """Read CSV with DictReader; return alumni-shaped dicts with data_source set."""
    base = os.path.basename(filepath)
    rows: list[dict[str, Any]] = []
    with open(filepath, newline="", encoding="utf-8-sig", errors="replace") as f:
        reader = csv.DictReader(f)
        for raw in reader:
            if not raw:
                continue
            alumni = _row_dict_to_alumni_fields(raw)
            alumni["data_source"] = base
            rows.append(alumni)
    return rows


def ingest_excel(filepath: str) -> list[dict[str, Any]]:
    """Read first sheet of Excel; first row is headers. Same normalisation as CSV."""
    from openpyxl import load_workbook

    base = os.path.basename(filepath)
    wb = load_workbook(filepath, read_only=True, data_only=True)
    try:
        ws = wb.active
        it = ws.iter_rows(values_only=True)
        header_row = next(it, None)
        if not header_row:
            return []
        headers = [_cell_str(h) for h in header_row]
        rows_out: list[dict[str, Any]] = []
        for row in it:
            if row is None:
                continue
            raw: dict[str, Any] = {}
            for i, h in enumerate(headers):
                if not h:
                    continue
                val = row[i] if i < len(row) else None
                raw[h] = val
            if not any(_cell_str(v) for v in raw.values()):
                continue
            alumni = _row_dict_to_alumni_fields(raw)
            alumni["data_source"] = base
            rows_out.append(alumni)
        return rows_out
    finally:
        wb.close()


def ingest_pdf(filepath: str) -> dict[str, Any]:
    from PyPDF2 import PdfReader

    base = os.path.basename(filepath)
    try:
        reader = PdfReader(filepath)
        parts: list[str] = []
        for page in reader.pages:
            try:
                t = page.extract_text()
                if t:
                    parts.append(t)
            except Exception as exc:  # noqa: BLE001
                logger.warning("PDF page extract failed in %s: %s", base, exc)
        text_content = "\n".join(parts)
        page_count = len(reader.pages)
        return {"filename": base, "text_content": text_content, "page_count": page_count}
    except Exception as exc:  # noqa: BLE001
        logger.warning("Failed to read PDF %s: %s", filepath, exc)
        return {"filename": base, "text_content": "", "page_count": 0}


def ingest_docx(filepath: str) -> dict[str, Any]:
    import docx

    base = os.path.basename(filepath)
    try:
        document = docx.Document(filepath)
        paras = [p.text for p in document.paragraphs]
        text_content = "\n".join(paras)
        paragraph_count = len(document.paragraphs)
        return {"filename": base, "text_content": text_content, "paragraph_count": paragraph_count}
    except Exception as exc:  # noqa: BLE001
        logger.warning("Failed to read Word document %s: %s", filepath, exc)
        return {"filename": base, "text_content": "", "paragraph_count": 0}


def ingest_txt(filepath: str) -> dict[str, Any]:
    """Plain text file; same unstructured shape with line_count."""
    base = os.path.basename(filepath)
    try:
        with open(filepath, encoding="utf-8", errors="replace") as f:
            text_content = f.read()
        line_count = len(text_content.splitlines())
        return {"filename": base, "text_content": text_content, "line_count": line_count}
    except Exception as exc:  # noqa: BLE001
        logger.warning("Failed to read text file %s: %s", filepath, exc)
        return {"filename": base, "text_content": "", "line_count": 0}


def ingest_all_files(folder_path: str = "data/sources/") -> dict[str, Any]:
    """
    Scan folder for ingestible files. Skips hidden entries (name starting with '.').
    Returns structured alumni rows, unstructured text blobs, and a summary string.
    """
    root = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", folder_path))
    structured: list[dict[str, Any]] = []
    unstructured: list[dict[str, Any]] = []

    n_csv = n_xlsx = n_pdf = n_docs = 0

    if not os.path.isdir(root):
        logger.warning("Folder does not exist: %s", root)
        summary = "Processed 0 CSVs, 0 Excel files, 0 PDFs, 0 documents"
        return {"structured": structured, "unstructured": unstructured, "summary": summary}

    for name in sorted(os.listdir(root)):
        if name.startswith("."):
            continue
        path = os.path.join(root, name)
        if not os.path.isfile(path):
            continue
        lower = name.lower()
        try:
            if lower.endswith(".csv"):
                structured.extend(ingest_csv(path))
                n_csv += 1
            elif lower.endswith(".xlsx"):
                structured.extend(ingest_excel(path))
                n_xlsx += 1
            elif lower.endswith(".pdf"):
                unstructured.append(ingest_pdf(path))
                n_pdf += 1
            elif lower.endswith(".docx"):
                unstructured.append(ingest_docx(path))
                n_docs += 1
            elif lower.endswith(".txt"):
                unstructured.append(ingest_txt(path))
                n_docs += 1
        except Exception as exc:  # noqa: BLE001
            logger.warning("Failed to ingest %s: %s", path, exc)

    summary = (
        f"Processed {n_csv} CSVs, {n_xlsx} Excel files, {n_pdf} PDFs, {n_docs} documents"
    )
    return {"structured": structured, "unstructured": unstructured, "summary": summary}


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    result = ingest_all_files()
    print(result["summary"])
